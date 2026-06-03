from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_table_margins(table, top=100, bottom=100, left=150, right=150):
    tblPr = table._tbl.tblPr
    tblCellMar = parse_xml(f'<w:tblCellMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tblCellMar>')
    tblPr.append(tblCellMar)

def set_cell_vertical_alignment(cell, align="center"):
    tcPr = cell._tc.get_or_add_tcPr()
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="{align}"/>')
    tcPr.append(vAlign)

def set_table_alignment(table, align=WD_TABLE_ALIGNMENT.CENTER):
    table.alignment = align

def set_column_widths(table, widths):
    for i, width in enumerate(widths):
        table.columns[i].width = width
        for row in table.rows:
            row.cells[i].width = width

def make_wrappable(text):
    if not text:
        return text
    for char in ['/', '?', '&', '=', ',', '|']:
        text = text.replace(char, char + '\u200b')
    return text

def add_page_number(run):
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve">PAGE</w:instrText>')
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def format_run(run, font_name="Arial", size_pt=11, bold=False, italic=False, color_rgb=None):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color_rgb:
        run.font.color.rgb = color_rgb

def add_heading_styled(doc, text, level, space_before=12, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    
    if level == 1:
        run = p.add_run(text)
        format_run(run, font_name="Arial", size_pt=16, bold=True, color_rgb=RGBColor(0x1B, 0x36, 0x5D))
    elif level == 2:
        run = p.add_run(text)
        format_run(run, font_name="Arial", size_pt=13, bold=True, color_rgb=RGBColor(0x2C, 0x52, 0x82))
    return p
