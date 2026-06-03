import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from config import (
    safe_print, 
    WEBDEFACEMENT_SIGNATURES, 
    CRYPTOJACKING_SIGNATURES, 
    REDIRECT_WHITELIST, 
    INFO_DISCLOSURE_PATHS
)

from .constants import (
    COLOR_PRIMARY_HEX,
    COLOR_LIGHT_GRAY_HEX,
    COLOR_ALERT_BG_HEX,
    COLOR_INACTIVE_BG_HEX,
    MONTHS_EN_ID,
    MITIGATIONS_LIST
)

from .helpers import (
    set_cell_background,
    set_table_margins,
    set_cell_vertical_alignment,
    set_table_alignment,
    set_column_widths,
    make_wrappable,
    add_page_number,
    format_run,
    add_heading_styled,
    add_bookmark,
    add_internal_link
)

def generate_docx_report(results):
    report_dir = os.path.dirname(os.path.abspath(__file__))
    base_dir = os.path.dirname(report_dir)
    screenshot_path = os.path.join(base_dir, "defaced_screenshot.png")
    output_path = os.path.join(base_dir, f"Laporan Hasil Scan Web Defacement - {datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}.docx")
    
    # Sort results so compromised domains are first, then active, then inactive
    results_sorted = sorted(results, key=lambda x: (
        not x.get("defaced", False), 
        not x.get("cryptojacking_detected", False), 
        not x.get("redirect_detected", False),
        not x.get("info_disclosure_detected", False),
        not x.get("ssl_error", False),
        not x.get("active", False), 
        x["url"]
    ))
    
    # Calculate stats
    total_domains = len(results)
    active_domains = sum(1 for x in results if x["active"])
    inactive_domains = total_domains - active_domains
    defaced_domains = sum(1 for x in results if x["defaced"])
    cryptojacking_domains = sum(1 for x in results if x.get("cryptojacking_detected", False))
    redirect_domains = sum(1 for x in results if x.get("redirect_detected", False))
    info_disclosure_domains = sum(1 for x in results if x.get("info_disclosure_detected", False))
    ssl_error_domains = sum(1 for x in results if x.get("ssl_error", False))
    
    current_time_str = datetime.now().strftime("%d %B %Y, %H:%M")
    current_year = datetime.now().strftime("%Y")
    for en, ind in MONTHS_EN_ID.items():
        if en in current_time_str:
            current_time_str = current_time_str.replace(en, ind)
            break
            
    doc = Document()
    
    # Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Configure page numbering in the footer (skipping the cover page)
    for section in doc.sections:
        section.different_first_page_header_footer = True
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run_text = p.add_run("Halaman ")
        format_run(run_text, font_name="Arial", size_pt=10, color_rgb=RGBColor(0x71, 0x80, 0x96))
        
        run_num = p.add_run()
        format_run(run_num, font_name="Arial", size_pt=10, color_rgb=RGBColor(0x71, 0x80, 0x96))
        add_page_number(run_num)

    # ------------------ COVER PAGE ------------------
    for _ in range(3):
        doc.add_paragraph()
        
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("LAPORAN HASIL PEMINDAIAN")
    format_run(r_title, font_name="Arial", size_pt=24, bold=True, color_rgb=RGBColor(0x1B, 0x36, 0x5D))
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(5)
    r_sub = p_sub.add_run("MONITORING KEAMANAN DAN INTEGRITAS WEBSITE\nPEMERINTAH KOTA PEKANBARU")
    format_run(r_sub, font_name="Arial", size_pt=14, bold=True, color_rgb=RGBColor(0x4A, 0x55, 0x68))

    p_sub_tag = doc.add_paragraph()
    p_sub_tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub_tag.paragraph_format.space_after = Pt(5)
    r_sub_tag = p_sub_tag.add_run("(*.pekanbaru.go.id)")
    format_run(r_sub_tag, font_name="Arial", size_pt=12, italic=True, color_rgb=RGBColor(0x71, 0x80, 0x96))
    
    p_sep = doc.add_paragraph()
    p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sep.paragraph_format.space_after = Pt(46)
    r_sep = p_sep.add_run("━" * 40)
    format_run(r_sep, font_name="Arial", size_pt=10, bold=True, color_rgb=RGBColor(0x71, 0x80, 0x96))
    
    vuln_status_parts = []
    if defaced_domains > 0:
        vuln_status_parts.append(f"{defaced_domains} URL Defacement/SEO")
    if cryptojacking_domains > 0:
        vuln_status_parts.append(f"{cryptojacking_domains} URL Cryptojacking")
    if redirect_domains > 0:
        vuln_status_parts.append(f"{redirect_domains} URL Redirects")
    if info_disclosure_domains > 0:
        vuln_status_parts.append(f"{info_disclosure_domains} URL Kebocoran Data")
    if ssl_error_domains > 0:
        vuln_status_parts.append(f"{ssl_error_domains} URL Error SSL/TLS")
        
    if vuln_status_parts:
        vuln_status_text = "🚨 Terindikasi: " + " & ".join(vuln_status_parts)
    else:
        vuln_status_text = "✅ Seluruh URL Bersih / Aman"

    metadata = [
        ("Tanggal Periksa", f"{current_time_str} WIB"),
        ("Diperiksa Oleh", "Sandiman (Analis & Operasi Cyber Security)"),
        ("Status Kerentanan", vuln_status_text),
        ("Jumlah URL yang Diperiksa", f"{total_domains} URL"),
    ]
    
    meta_table = doc.add_table(rows=len(metadata), cols=2)
    set_table_alignment(meta_table)
    set_table_margins(meta_table, top=80, bottom=80, left=120, right=120)
    
    for idx, (label, val) in enumerate(metadata):
        row = meta_table.rows[idx]
        
        c0 = row.cells[0]
        c0.text = ""
        r0 = c0.paragraphs[0].add_run(label)
        format_run(r0, font_name="Arial", size_pt=11, bold=True, color_rgb=RGBColor(0x2D, 0x37, 0x48))
        set_cell_background(c0, COLOR_LIGHT_GRAY_HEX)
        
        c1 = row.cells[1]
        c1.text = ""
        r1 = c1.paragraphs[0].add_run(val)
        if "🚨" in val:
            format_run(r1, font_name="Arial", size_pt=11, bold=True, color_rgb=RGBColor(0x99, 0x1B, 0x1B))
        elif "✅" in val:
            format_run(r1, font_name="Arial", size_pt=11, bold=True, color_rgb=RGBColor(0x16, 0x65, 0x34))
        else:
            format_run(r1, font_name="Arial", size_pt=11, color_rgb=RGBColor(0x2D, 0x37, 0x48)) 
            
    set_column_widths(meta_table, [Inches(3.0), Inches(4.5)])
    
    for _ in range(6):
        doc.add_paragraph()
        
    p_org = doc.add_paragraph()
    p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_org = p_org.add_run(f"DINAS KOMUNIKASI INFORMATIKA STATISTIK DAN PERSANDIAN\nKOTA PEKANBARU\n{current_year}")
    format_run(r_org, font_name="Arial", size_pt=10, bold=True, color_rgb=RGBColor(0x71, 0x80, 0x96))
    
    doc.add_page_break()

    # ------------------ SECTION 1: RINGKASAN EKSEKUTIF ------------------
    add_heading_styled(doc, "1. Ringkasan Eksekutif", level=1)
    
    p_desc1 = doc.add_paragraph()
    p_desc1.paragraph_format.space_after = Pt(8)
    p_desc1.paragraph_format.line_spacing = 1.15
    p_desc1.paragraph_format.first_line_indent = Inches(0.5)
    r_desc1 = p_desc1.add_run(
        f"Pada tanggal {current_time_str.split(',')[0]}, telah dilaksanakan pemindaian rutin terhadap domain/subdomain milik Pemerintah Kota Pekanbaru (*.pekanbaru.go.id). "
        "Pemindaian ini secara khusus menargetkan pendeteksian serangan Web Defacement, SEO Poisoning (terkait judi online), Cryptojacking (injeksi skrip penambang aset kripto), serta Malicious Redirects & Traffic Hijacking (pengalihan lalu lintas ilegal). "
        "Metode ini mendeteksi keberadaan kata kunci terlarang judi online, skrip penambangan, konfigurasi pengalihan HTTP, dan pengalihan berbasis klien (HTML meta & JS location) yang disusupkan pada situs web instansi daerah."
    )
    p_desc1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    format_run(r_desc1, font_name="Arial", size_pt=11, color_rgb=RGBColor(0x2D, 0x37, 0x48))

    p_desc2 = doc.add_paragraph()
    p_desc2.paragraph_format.space_after = Pt(12)
    p_desc2.paragraph_format.line_spacing = 1.15
    p_desc2.paragraph_format.first_line_indent = Inches(0.5)
    p_desc2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    affected_narratives = []
    if defaced_domains > 0:
        defaced_names = ", ".join([x["url"].replace("https://", "").replace("http://", "") for x in results if x["defaced"]])
        affected_narratives.append(f"terdapat {defaced_domains} URL aktif yang terkonfirmasi mengalami insiden SEO Poisoning")
    if cryptojacking_domains > 0:
        cj_names = ", ".join([x["url"].replace("https://", "").replace("http://", "") for x in results if x.get("cryptojacking_detected", False)])
        affected_narratives.append(f"terdapat {cryptojacking_domains} URL aktif yang terkonfirmasi mengalami insiden Cryptojacking")
    if redirect_domains > 0:
        redir_names = ", ".join([x["url"].replace("https://", "").replace("http://", "") for x in results if x.get("redirect_detected", False)])
        affected_narratives.append(f"terdapat {redirect_domains} URL aktif yang terkonfirmasi mengalami insiden Malicious Redirect")
    if info_disclosure_domains > 0:
        id_names = ", ".join([x["url"].replace("https://", "").replace("http://", "") for x in results if x.get("info_disclosure_detected", False)])
        affected_narratives.append(f"terdapat {info_disclosure_domains} URL aktif yang terkonfirmasi mengalami kebocoran data sensitif / Information Disclosure")
    if ssl_error_domains > 0:
        ssl_names = ", ".join([x["url"].replace("https://", "").replace("http://", "") for x in results if x.get("ssl_error", False)])
        affected_narratives.append(f"terdapat {ssl_error_domains} URL yang terdeteksi mengalami kesalahan sertifikat SSL/TLS")

    if affected_narratives:
        status_narative = (
            f"Dari total {total_domains} URL yang didaftarkan dalam daftar pantauan, hasil pemindaian menunjukkan bahwa "
            f"{' dan '.join(affected_narratives)}. "
            "Aktivitas ini dapat menurunkan reputasi domain pemerintah (.go.id), membebani kinerja server dan perangkat pengguna, serta merusak indeks pencarian di mesin pencari. "
            f"Sementara itu, sebanyak {active_domains - sum(1 for x in results if x['defaced'] or x.get('cryptojacking_detected') or x.get('redirect_detected') or x.get('info_disclosure_detected') or x.get('ssl_error'))} URL aktif lainnya dinyatakan AMAN, "
            f"dan {inactive_domains} URL dalam kondisi NONAKTIF atau mengalami error koneksi pada saat pemindaian dilakukan."
        )
    else:
        status_narative = (
            f"Dari total {total_domains} URL yang diperiksa, hasil pemindaian menunjukkan seluruh URL dalam kondisi AMAN dan "
            "tidak ditemukan indikasi serangan Web Defacement, SEO Poisoning, penyusupan skrip cryptojacking, malicious redirect, maupun kebocoran data pada halaman utama. "
            f"Sebanyak {active_domains} URL aktif berjalan normal, dan {inactive_domains} URL berada dalam kondisi NONAKTIF."
        )
        
    r_desc2 = p_desc2.add_run(status_narative)
    format_run(r_desc2, font_name="Arial", size_pt=11, color_rgb=RGBColor(0x2D, 0x37, 0x48))

    add_heading_styled(doc, "Statistik Hasil Pemindaian", level=2)
    stats_table = doc.add_table(rows=9, cols=2)
    set_table_alignment(stats_table)
    set_table_margins(stats_table, top=100, bottom=100, left=150, right=150)
    
    stats_data = [
        ("Parameter Pemindaian", "Nilai / Jumlah"),
        ("Total URL Diperiksa", str(total_domains)),
        ("URL Status Aktif (200 OK)", str(active_domains)),
        ("URL Status Nonaktif / Timeout", str(inactive_domains)),
        ("URL Terindikasi Defacement / SEO Poisoning", str(defaced_domains)),
        ("URL Terindikasi Cryptojacking", str(cryptojacking_domains)),
        ("URL Terindikasi Malicious Redirect", str(redirect_domains)),
        ("URL Terindikasi Kebocoran Data\n(Information Disclosure)", str(info_disclosure_domains)),
        ("URL Terindikasi Error SSL/TLS", str(ssl_error_domains))
    ]
    
    for idx, (param, val) in enumerate(stats_data):
        row = stats_table.rows[idx]
        
        c0 = row.cells[0]
        c0.text = ""
        r0 = c0.paragraphs[0].add_run(param)
        
        c1 = row.cells[1]
        c1.text = ""
        
        if idx == 0:
            r1 = c1.paragraphs[0].add_run(val)
            format_run(r0, font_name="Arial", size_pt=11, bold=True, color_rgb=RGBColor(255, 255, 255))
            format_run(r1, font_name="Arial", size_pt=11, bold=True, color_rgb=RGBColor(255, 255, 255))
            set_cell_background(c0, COLOR_PRIMARY_HEX)
            set_cell_background(c1, COLOR_PRIMARY_HEX)
        else:
            format_run(r0, font_name="Arial", size_pt=11, color_rgb=RGBColor(0x2D, 0x37, 0x48))
            
            # Highlight row for defacement
            if idx == 4 and defaced_domains > 0:
                p_count = c1.paragraphs[0]
                r_count = p_count.add_run(val)
                format_run(r_count, font_name="Arial", size_pt=11, bold=True, color_rgb=RGBColor(0x99, 0x1B, 0x1B))
                
                defaced_names = ", ".join([x["url"].replace("https://", "").replace("http://", "") for x in results if x["defaced"]])
                for name in defaced_names.split(", "):
                    p_bullet = c1.add_paragraph(style='List Bullet')
                    p_bullet.paragraph_format.space_before = Pt(1)
                    p_bullet.paragraph_format.space_after = Pt(1)
                    p_bullet.paragraph_format.left_indent = Inches(0.2)
                    r_bullet = p_bullet.add_run(name)
                    format_run(r_bullet, font_name="Arial", size_pt=9.5, italic=True, color_rgb=RGBColor(0x99, 0x1B, 0x1B))
            
            # Highlight row for cryptojacking
            elif idx == 5 and cryptojacking_domains > 0:
                p_count = c1.paragraphs[0]
                r_count = p_count.add_run(val)
                format_run(r_count, font_name="Arial", size_pt=11, bold=True, color_rgb=RGBColor(0x99, 0x1B, 0x1B))
                
                cj_names = ", ".join([x["url"].replace("https://", "").replace("http://", "") for x in results if x.get("cryptojacking_detected", False)])
                for name in cj_names.split(", "):
                    p_bullet = c1.add_paragraph(style='List Bullet')
                    p_bullet.paragraph_format.space_before = Pt(1)
                    p_bullet.paragraph_format.space_after = Pt(1)
                    p_bullet.paragraph_format.left_indent = Inches(0.2)
                    r_bullet = p_bullet.add_run(name)
                    format_run(r_bullet, font_name="Arial", size_pt=9.5, italic=True, color_rgb=RGBColor(0x99, 0x1B, 0x1B))
            
            # Highlight row for redirect
            elif idx == 6 and redirect_domains > 0:
                p_count = c1.paragraphs[0]
                r_count = p_count.add_run(val)
                format_run(r_count, font_name="Arial", size_pt=11, bold=True, color_rgb=RGBColor(0x99, 0x1B, 0x1B))
                
                redir_names = ", ".join([x["url"].replace("https://", "").replace("http://", "") for x in results if x.get("redirect_detected", False)])
                for name in redir_names.split(", "):
                    p_bullet = c1.add_paragraph(style='List Bullet')
                    p_bullet.paragraph_format.space_before = Pt(1)  
                    p_bullet.paragraph_format.space_after = Pt(1)
                    p_bullet.paragraph_format.left_indent = Inches(0.2)
                    r_bullet = p_bullet.add_run(name)
                    format_run(r_bullet, font_name="Arial", size_pt=9.5, italic=True, color_rgb=RGBColor(0x99, 0x1B, 0x1B))
            
            # Highlight row for info disclosure
            elif idx == 7 and info_disclosure_domains > 0:
                p_count = c1.paragraphs[0]
                r_count = p_count.add_run(val)
                format_run(r_count, font_name="Arial", size_pt=11, bold=True, color_rgb=RGBColor(0x99, 0x1B, 0x1B))
                
                id_names = ", ".join([x["url"].replace("https://", "").replace("http://", "") for x in results if x.get("info_disclosure_detected", False)])
                for name in id_names.split(", "):
                    p_bullet = c1.add_paragraph(style='List Bullet')
                    p_bullet.paragraph_format.space_before = Pt(1)
                    p_bullet.paragraph_format.space_after = Pt(1)
                    p_bullet.paragraph_format.left_indent = Inches(0.2)
                    r_bullet = p_bullet.add_run(name)
                    format_run(r_bullet, font_name="Arial", size_pt=9.5, italic=True, color_rgb=RGBColor(0x99, 0x1B, 0x1B))
                    
            # Highlight row for SSL/TLS errors
            elif idx == 8 and ssl_error_domains > 0:
                p_count = c1.paragraphs[0]
                r_count = p_count.add_run(val)
                format_run(r_count, font_name="Arial", size_pt=11, bold=True, color_rgb=RGBColor(0x99, 0x1B, 0x1B))
                
                ssl_names = ", ".join([x["url"].replace("https://", "").replace("http://", "") for x in results if x.get("ssl_error", False)])
                for name in ssl_names.split(", "):
                    p_bullet = c1.add_paragraph(style='List Bullet')
                    p_bullet.paragraph_format.space_before = Pt(1)
                    p_bullet.paragraph_format.space_after = Pt(1)
                    p_bullet.paragraph_format.left_indent = Inches(0.2)
                    r_bullet = p_bullet.add_run(name)
                    format_run(r_bullet, font_name="Arial", size_pt=9.5, italic=True, color_rgb=RGBColor(0x99, 0x1B, 0x1B))
            else:
                r1 = c1.paragraphs[0].add_run(val)
                is_alert = (idx == 3 and inactive_domains > 0) or (idx == 4 and defaced_domains > 0) or (idx == 5 and cryptojacking_domains > 0) or (idx == 6 and redirect_domains > 0) or (idx == 7 and info_disclosure_domains > 0) or (idx == 8 and ssl_error_domains > 0)
                format_run(r1, font_name="Arial", size_pt=11, bold=True if idx in [3, 4, 5, 6, 7, 8] else False, 
                           color_rgb=RGBColor(0x99, 0x1B, 0x1B) if is_alert else RGBColor(0x2D, 0x37, 0x48))
            
            bg = COLOR_LIGHT_GRAY_HEX if idx % 2 == 1 else "FFFFFF"
            if idx in [3, 4, 5, 6, 7, 8] and (
                (idx == 3 and inactive_domains > 0) or
                (idx == 4 and defaced_domains > 0) or 
                (idx == 5 and cryptojacking_domains > 0) or 
                (idx == 6 and redirect_domains > 0) or 
                (idx == 7 and info_disclosure_domains > 0) or 
                (idx == 8 and ssl_error_domains > 0)
            ):
                bg = COLOR_ALERT_BG_HEX
            set_cell_background(c0, bg)
            set_cell_background(c1, bg)
            
    set_column_widths(stats_table, [Inches(4.5), Inches(2.5)])
    
    doc.add_paragraph()

    # ------------------ SECTION 2: METODOLOGI ------------------
    add_heading_styled(doc, "2. Metodologi Pemindaian", level=1)
    
    p_meth = doc.add_paragraph()
    p_meth.paragraph_format.space_after = Pt(8)
    p_meth.paragraph_format.line_spacing = 1.15
    p_meth.paragraph_format.first_line_indent = Inches(0.5)
    p_meth.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_meth = p_meth.add_run(
        "Pemindaian dilakukan secara otomatis menggunakan skrip keamanan yang mengintegrasikan library "
        "requests untuk melakukan koneksi HTTP/HTTPS dan BeautifulSoup4 untuk menganalisis dokumen HTML yang dihasilkan. "
        "Adapun alur deteksi yang diterapkan adalah sebagai berikut:"
    )
    format_run(r_meth, font_name="Arial", size_pt=11, color_rgb=RGBColor(0x2D, 0x37, 0x48))
    
    steps = [
        "Pengambilan data daftar target URL aktif dari file listsubdomain.txt.",
        "Pengiriman HTTP GET Request dengan custom User-Agent modern untuk menyimulasikan akses peramban (browser) normal guna menghindari pemblokiran WAF.",
        "Pengecekan HTTP Response Status Code: jika mengembalikan nilai 200 OK, situs dianalisis lebih lanjut (Aktif). Jika tidak atau gagal koneksi, dikategorikan Nonaktif.",
        "Evaluasi riwayat status HTTP (Response History) untuk mendeteksi status code 301/302 yang mengarah ke luar whitelist.",
        "Pencarian konten teks dan tag jangkar (anchor tags) secara sensitif terhadap kata kunci terlarang (blacklist keywords) judi online.",
        "Pencarian pada seluruh source code HTML (termasuk tag <script>) untuk menemukan signature, domain, atau pustaka JavaScript cryptojacking (penambang kripto).",
        "Ekstraksi tag HTML <meta http-equiv=\"refresh\"> dan kode JavaScript client-side (window.location / location.replace) untuk mendeteksi pengalihan tersembunyi.",
        "Jika ditemukan indikasi defacement/SEO poisoning, cryptojacking, atau malicious redirect pada respon halaman aktif, sistem menandai URL tersebut dan mengkategorikan kerentanannya."
    ]
    
    for s in steps:
        p_step = doc.add_paragraph(style='List Bullet')
        p_step.paragraph_format.space_after = Pt(4)
        p_step.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r_step = p_step.add_run(s)
        format_run(r_step, font_name="Arial", size_pt=11, color_rgb=RGBColor(0x2D, 0x37, 0x48))

    p_kw_title = doc.add_paragraph()
    p_kw_title.paragraph_format.space_before = Pt(8)
    p_kw_title.paragraph_format.space_after = Pt(4)
    r_kw_title = p_kw_title.add_run("Kata Kunci Blacklist Judi Online:")
    format_run(r_kw_title, font_name="Arial", size_pt=11, bold=True, color_rgb=RGBColor(0x2D, 0x37, 0x48))
    
    p_kw = doc.add_paragraph()
    p_kw.paragraph_format.space_after = Pt(8)
    r_kw = p_kw.add_run(", ".join(WEBDEFACEMENT_SIGNATURES))
    format_run(r_kw, font_name="Arial", size_pt=11, italic=True, color_rgb=RGBColor(0x4A, 0x55, 0x68))

    p_cj_title = doc.add_paragraph()
    p_cj_title.paragraph_format.space_before = Pt(8)
    p_cj_title.paragraph_format.space_after = Pt(4)
    r_cj_title = p_cj_title.add_run("Signature Cryptojacking yang Dipantau:")
    format_run(r_cj_title, font_name="Arial", size_pt=11, bold=True, color_rgb=RGBColor(0x2D, 0x37, 0x48))
    
    p_cj = doc.add_paragraph()
    p_cj.paragraph_format.space_after = Pt(12)
    r_cj = p_cj.add_run(", ".join(CRYPTOJACKING_SIGNATURES))
    format_run(r_cj, font_name="Arial", size_pt=11, italic=True, color_rgb=RGBColor(0x4A, 0x55, 0x68))

    p_id_title = doc.add_paragraph()
    p_id_title.paragraph_format.space_before = Pt(8)
    p_id_title.paragraph_format.space_after = Pt(4)
    r_id_title = p_id_title.add_run("Path Path Information Disclosure yang Dipantau:")
    format_run(r_id_title, font_name="Arial", size_pt=11, bold=True, color_rgb=RGBColor(0x2D, 0x37, 0x48))
    
    p_id = doc.add_paragraph()
    p_id.paragraph_format.space_after = Pt(12)
    r_id = p_id.add_run(", ".join([path for path, _ in INFO_DISCLOSURE_PATHS]))
    format_run(r_id, font_name="Arial", size_pt=11, italic=True, color_rgb=RGBColor(0x4A, 0x55, 0x68))

    # ------------------ SECTION 3: HASIL PEMINDAIAN SUBDOMAIN ------------------
    add_heading_styled(doc, "3. Hasil Pemindaian", level=1)
    
    p_table_desc = doc.add_paragraph()
    p_table_desc.paragraph_format.space_after = Pt(8)
    p_table_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_table_desc.paragraph_format.first_line_indent = Inches(0.5)
    r_table_desc = p_table_desc.add_run(
        f"Berikut merupakan tabel lengkap hasil pemindaian ke-{total_domains} URL Pemerintah Kota Pekanbaru yang dilakukan pada {current_time_str.split(',')[0]}. "
        "Baris berwarna merah menandakan ditemukannya kerentanan, "
        "baris berwarna hijau terang menandakan situs aman, dan baris abu-abu menandakan status nonaktif/gagal akses."
    )
    format_run(r_table_desc, font_name="Arial", size_pt=11, color_rgb=RGBColor(0x2D, 0x37, 0x48))

    results_table = doc.add_table(rows=1, cols=8)
    set_table_alignment(results_table)
    set_table_margins(results_table, top=60, bottom=60, left=80, right=80)
    
    hdr_cells = results_table.rows[0].cells
    headers_text = ["No", "Domain", "Status\nAktif", "Defacement", "Crypto\njacking", "Redirect", "Kebocoran\nData", "Rujukan"]
    for i, title in enumerate(headers_text):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        format_run(r, font_name="Arial", size_pt=8.5, bold=True, color_rgb=RGBColor(255, 255, 255))
        set_cell_background(hdr_cells[i], COLOR_PRIMARY_HEX)
        set_cell_vertical_alignment(hdr_cells[i], "center")
        
    set_column_widths(results_table, [Inches(0.35), Inches(1.45), Inches(0.85), Inches(0.8), Inches(0.8), Inches(0.8), Inches(0.8), Inches(0.7)])
    
    affected_count = 0
    for idx, item in enumerate(results_sorted, start=1):
        row = results_table.add_row()
        cells = row.cells
        
        url = item["url"]
        active = item["active"]
        defaced = item["defaced"]
        cryptojacking = item.get("cryptojacking_detected", False)
        redirect = item.get("redirect_detected", False)
        info_disclosure = item.get("info_disclosure_detected", False)
        ssl_error = item.get("ssl_error", False)
        
        if ssl_error:
            status_aktif_txt = "Aktif (SSL Err)" if active else "SSL Error"
        else:
            status_aktif_txt = "Aktif" if active else "Nonaktif"
        
        is_vuln = (defaced or cryptojacking or redirect or info_disclosure)
        if is_vuln or ssl_error:
            if is_vuln:
                affected_count += 1
                rujukan_txt = f"Lampiran {affected_count}"
            else:
                rujukan_txt = ""
            defacement_txt = "❌" if defaced else "✔"
            cj_txt = "❌" if cryptojacking else "✔"
            redir_txt = "❌" if redirect else "✔"
            id_txt = "❌" if info_disclosure else "✔"
            row_bg = COLOR_ALERT_BG_HEX
            txt_color = RGBColor(0x99, 0x1B, 0x1B)
            cell_bold = True
        elif not active:
            defacement_txt = "-"
            cj_txt = "-"
            redir_txt = "-"
            id_txt = "-"
            rujukan_txt = ""
            row_bg = COLOR_INACTIVE_BG_HEX
            txt_color = RGBColor(0x64, 0x74, 0x8B)
            cell_bold = False
        else:
            defacement_txt = "✔"
            cj_txt = "✔"
            redir_txt = "✔"
            id_txt = "✔"
            rujukan_txt = ""
            row_bg = "FFFFFF" if idx % 2 == 0 else COLOR_LIGHT_GRAY_HEX
            txt_color = RGBColor(0x2D, 0x37, 0x48)
            cell_bold = False
            
        row_data = [str(idx), url, status_aktif_txt, defacement_txt, cj_txt, redir_txt, id_txt, rujukan_txt]
        
        for i, text in enumerate(row_data):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            if i in [0, 2, 3, 4, 5, 6, 7]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            if i == 7 and rujukan_txt:
                bookmark_name = f"lampiran_{affected_count}"
                add_internal_link(p, bookmark_name, text)
            else:
                r = p.add_run(text)
                
                # Formatting based on type
                if i == 2: # Status Aktif
                    if "SSL" in text or "SSL Err" in text or "SSL Error" in text:
                        format_run(r, font_name="Arial", size_pt=8, bold=True, color_rgb=RGBColor(0x99, 0x1B, 0x1B))
                    else:
                        format_run(r, font_name="Arial", size_pt=8, bold=cell_bold, color_rgb=txt_color)
                elif i == 3: # Web Defacement
                    if defaced:
                        format_run(r, font_name="Arial", size_pt=8, bold=True, color_rgb=RGBColor(0x99, 0x1B, 0x1B))
                    elif active:
                        format_run(r, font_name="Arial", size_pt=8, bold=True, color_rgb=RGBColor(0x16, 0x65, 0x34))
                    else:
                        format_run(r, font_name="Arial", size_pt=8, color_rgb=txt_color)
                elif i == 4: # Cryptojacking
                    if cryptojacking:
                        format_run(r, font_name="Arial", size_pt=8, bold=True, color_rgb=RGBColor(0x99, 0x1B, 0x1B))
                    elif active:
                        format_run(r, font_name="Arial", size_pt=8, bold=True, color_rgb=RGBColor(0x16, 0x65, 0x34))
                    else:
                        format_run(r, font_name="Arial", size_pt=8, color_rgb=txt_color)
                elif i == 5: # Redirect
                    if redirect:
                        format_run(r, font_name="Arial", size_pt=8, bold=True, color_rgb=RGBColor(0x99, 0x1B, 0x1B))
                    elif active:
                        format_run(r, font_name="Arial", size_pt=8, bold=True, color_rgb=RGBColor(0x16, 0x65, 0x34))
                    else:
                        format_run(r, font_name="Arial", size_pt=8, color_rgb=txt_color)
                elif i == 6: # Information Disclosure
                    if info_disclosure:
                        format_run(r, font_name="Arial", size_pt=8, bold=True, color_rgb=RGBColor(0x99, 0x1B, 0x1B))
                    elif active:
                        format_run(r, font_name="Arial", size_pt=8, bold=True, color_rgb=RGBColor(0x16, 0x65, 0x34))
                    else:
                        format_run(r, font_name="Arial", size_pt=8, color_rgb=txt_color)
                else:
                    format_run(r, font_name="Arial", size_pt=8, bold=cell_bold, color_rgb=txt_color)
                
            set_cell_background(cells[i], row_bg)
            set_cell_vertical_alignment(cells[i], "center")
            
    set_column_widths(results_table, [Inches(0.35), Inches(1.45), Inches(0.85), Inches(0.8), Inches(0.8), Inches(0.8), Inches(0.8), Inches(0.7)])
    
    doc.add_paragraph()

    # ------------------ SECTION 4: DETAIL TEMUAN ------------------
    affected_items = [x for x in results_sorted if x["defaced"] or x.get("cryptojacking_detected", False) or x.get("redirect_detected", False) or x.get("info_disclosure_detected", False)]
    if affected_items:
        add_heading_styled(doc, "4. Detail Temuan Kerentanan", level=1)
        p_det_desc = doc.add_paragraph()
        p_det_desc.paragraph_format.space_after = Pt(12)
        p_det_desc.paragraph_format.line_spacing = 1.15
        p_det_desc.paragraph_format.first_line_indent = Inches(0.5)
        p_det_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r_det_desc = p_det_desc.add_run(
            "Berikut merupakan rincian temuan keamanan untuk URL yang terindikasi mengalami insiden keamanan Web Defacement / SEO Poisoning, Cryptojacking, Malicious Redirect, atau Kebocoran Data Sensitif:"
        )
        format_run(r_det_desc, font_name="Arial", size_pt=11, color_rgb=RGBColor(0x2D, 0x37, 0x48))
        
        for item_idx, item in enumerate(affected_items, start=1):
            p_heading = add_heading_styled(doc, f"Lampiran {item_idx}: {item['url'].replace('https://','').replace('http://','')}", level=2)
            add_bookmark(p_heading, f"lampiran_{item_idx}")
            
            detail_table = doc.add_table(rows=6, cols=2)
            set_table_alignment(detail_table)
            detail_table.allow_autofit = False
            set_table_margins(detail_table, top=100, bottom=100, left=150, right=150)
            
            categories = []
            findings = []
            impacts = []
            descriptions = []
            
            if item["defaced"]:
                categories.append("SEO Poisoning / Injeksi Tautan Judi Online")
                findings.append(f"Keywords: {', '.join(item['keywords_found'])}")
                impacts.append("Kerusakan reputasi domain pemerintah (.go.id) di mesin pencari (SEO), potensi deindexing dari Google Search, dan penyebaran konten ilegal.")
                descriptions.append("Halaman utama situs web disusupi oleh tautan jangkar (anchor links) tidak terlihat/tersembunyi yang mengarah ke berbagai situs eksternal judi online.")
                
            if item.get("cryptojacking_detected", False):
                categories.append("Cryptojacking / Injeksi Skrip Penambangan Kripto")
                findings.append(f"Signatures: {', '.join(item['cryptojacking_signatures'])}")
                impacts.append("Penyalahgunaan sumber daya CPU pengunjung (High CPU Usage), penurunan performa sistem, serta berkurangnya kenyamanan dan keamanan pengguna.")
                descriptions.append("Halaman utama situs web disusupi skrip penambangan cryptocurrency (seperti CoinHive/CryptoLoot) yang berjalan otomatis di browser pengunjung.")
                
            if item.get("redirect_detected", False):
                categories.append("Malicious Redirect & Traffic Hijacking")
                findings.append(f"Redirects: {', '.join(item['redirect_details'])}")
                impacts.append("Pengunjung situs dialihkan paksa ke situs berbahaya (phishing/malware), merusak reputasi instansi, dan deindexing oleh mesin pencari.")
                descriptions.append("Situs menyalahgunakan HTTP status redirect atau disusupi elemen pengalihan DOM client-side (meta refresh/skrip pengalihan location) ke domain tidak dikenal di luar whitelist.")

            if item.get("info_disclosure_detected", False):
                categories.append("Information Disclosure / Kebocoran Data Sensitif")
                findings.append(f"Exposed: {', '.join(item['info_disclosure_details'])}")
                impacts.append("Kebocoran informasi server, variabel kredensial (database/API keys), struktur repositori kode internal, serta detail konfigurasi PHP yang mempermudah penyerangan lanjutan.")
                descriptions.append("Server membiarkan file konfigurasi sensitif (seperti berkas .env, direktori .git/config, atau halaman pengujian phpinfo.php) dapat diakses secara publik karena kesalahan izin konfigurasi web server.")

            cat_str = " & ".join(categories)
            findings_str = " | ".join(findings)
            impact_str = " ".join(impacts)
            desc_str = " ".join(descriptions)
            
            detail_data = [
                ("Nama URL", make_wrappable(item["url"])),
                ("Status Koneksi", f"Aktif (HTTP {item['status_code']})" if item['active'] else "Nonaktif (SSL Error / Gagal Koneksi)"),
                ("Kategori Insiden", cat_str),
                ("Bukti Temuan", make_wrappable(findings_str)),
                ("Dampak Temuan", impact_str),
                ("Deskripsi Teknis", desc_str)
            ]
            
            for idx, (label, val) in enumerate(detail_data):
                row = detail_table.rows[idx]
                c0 = row.cells[0]
                c0.width = Inches(2.2)
                c0.text = ""
                r0 = c0.paragraphs[0].add_run(label)
                format_run(r0, font_name="Arial", size_pt=10, bold=True, color_rgb=RGBColor(0x2D, 0x37, 0x48))
                set_cell_background(c0, COLOR_LIGHT_GRAY_HEX)
                
                c1 = row.cells[1]
                c1.width = Inches(4.3)
                c1.text = ""
                r1 = c1.paragraphs[0].add_run(val)
                
                if idx == 0:
                    format_run(r1, font_name="Arial", size_pt=10, bold=True, color_rgb=RGBColor(0x1B, 0x36, 0x5D))
                elif idx == 3:
                    format_run(r1, font_name="Arial", size_pt=10, bold=True, color_rgb=RGBColor(0x99, 0x1B, 0x1B))
                else:
                    format_run(r1, font_name="Arial", size_pt=10, color_rgb=RGBColor(0x2D, 0x37, 0x48))
                    
                set_cell_background(c1, "FFFFFF")
                
            set_column_widths(detail_table, [Inches(2.2), Inches(4.3)])
            
            # Show visual proof if it's the defaced one and has screenshot
            if item["defaced"]:
                doc.add_paragraph()
                add_heading_styled(doc, "Bukti Visual Tangkapan Layar (Screenshot)", level=2)
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_before = Pt(6)
                p_img.paragraph_format.space_after = Pt(6)
                
                if os.path.exists(screenshot_path):
                    p_img.add_run().add_picture(screenshot_path, width=Inches(5.8))
                    p_cap = doc.add_paragraph()
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_cap.paragraph_format.space_after = Pt(18)
                    r_cap = p_cap.add_run(f"Gambar 1.{item_idx}: Tangkapan layar bukti injeksi kata kunci judi online pada {item['url']} - {current_time_str}")
                    format_run(r_cap, font_name="Arial", size_pt=9.5, italic=True, color_rgb=RGBColor(0x71, 0x80, 0x96))
                else:
                    r_img = p_img.add_run("[Berkas screenshot 'defaced_screenshot.png' tidak ditemukan untuk disematkan]")
                    format_run(r_img, font_name="Arial", size_pt=11, italic=True, color_rgb=RGBColor(0x99, 0x1B, 0x1B))
            
            doc.add_paragraph()

    # ------------------ SECTION 5: MITIGASI ------------------
    has_incident = (defaced_domains > 0 or cryptojacking_domains > 0 or redirect_domains > 0 or info_disclosure_domains > 0 or ssl_error_domains > 0)
    add_heading_styled(doc, "5. Rekomendasi Tindakan Mitigasi" if has_incident else "4. Rekomendasi Tindakan Mitigasi", level=1)
    
    p_mit_intro = doc.add_paragraph()
    p_mit_intro.paragraph_format.space_after = Pt(8)
    p_mit_intro.paragraph_format.line_spacing = 1.15
    p_mit_intro.paragraph_format.first_line_indent = Inches(0.5)
    p_mit_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    if has_incident:
        mit_intro_txt = (
            "Guna menangani insiden keamanan serta mencegah serangan serupa di masa mendatang, "
            "pihak administrator situs web terdampak direkomendasikan untuk segera melakukan langkah-langkah mitigasi berikut:"
        )
    else:
        mit_intro_txt = (
            "Meskipun seluruh URL saat ini dalam kondisi AMAN, langkah-langkah penguatan keamanan (hardening) proaktif berikut "
            "sangat direkomendasikan untuk mencegah insiden Web Defacement, SEO Poisoning, Cryptojacking, Malicious Redirect, Kebocoran Data Sensitif, dan Masalah SSL/TLS di masa mendatang:"
        )
    r_mit_intro = p_mit_intro.add_run(mit_intro_txt)
    format_run(r_mit_intro, font_name="Arial", size_pt=11, color_rgb=RGBColor(0x2D, 0x37, 0x48))
    
    for title, desc in MITIGATIONS_LIST:
        p_title = doc.add_paragraph()
        p_title.paragraph_format.space_before = Pt(6)
        p_title.paragraph_format.space_after = Pt(2)
        p_title.paragraph_format.keep_with_next = True
        r_t = p_title.add_run(title)
        format_run(r_t, font_name="Arial", size_pt=11, bold=True, color_rgb=RGBColor(0x1B, 0x36, 0x5D))
        
        p_d = doc.add_paragraph()
        p_d.paragraph_format.space_after = Pt(6)
        p_d.paragraph_format.line_spacing = 1.15
        p_d.paragraph_format.left_indent = Inches(0.2)
        p_d.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r_d = p_d.add_run(desc)
        format_run(r_d, font_name="Arial", size_pt=11, color_rgb=RGBColor(0x2D, 0x37, 0x48))
        
    doc.add_paragraph()
    p_sig = doc.add_paragraph()
    p_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_sig.paragraph_format.space_before = Pt(24)
    r_sig = p_sig.add_run(
        f"Pekanbaru, {current_time_str.split(',')[0]}\n\n\n\n"
        "___________________________\n"
        "Dinas Komunikasi Informatika Statistik dan Persandian\n"
        "Pemerintah Kota Pekanbaru"
    )
    format_run(r_sig, font_name="Arial", size_pt=10, bold=True, color_rgb=RGBColor(0x4A, 0x55, 0x68))

    try:
        doc.save(output_path)
        safe_print(f"Menyusun laporan hasil scan ke: {output_path}")
        safe_print("Laporan hasil pemeriksaan (.docx) berhasil disimpan!")
    except PermissionError:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        dir_name = os.path.dirname(output_path)
        base_name = os.path.basename(output_path).replace(".docx", "")
        fallback_path = os.path.join(dir_name, f"{base_name} ({timestamp}).docx")
        doc.save(fallback_path)
        safe_print(f"\n⚠️  [PERINGATAN] Berkas {output_path} sedang terbuka atau terkunci.")
        safe_print(f"Laporan dialihkan dan disimpan ke: {fallback_path}")
